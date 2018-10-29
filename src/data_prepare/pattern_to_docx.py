# -*- coding: utf-8 -*-
"""
Created on Thu Jul 30 16:21:28 2018

@author: bob.lee
"""
from win32com import client
import os
import time
import subprocess
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
from src.data_prepare.zip_decode import zip_main
import docx

DATA_HOME = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir, 'temp/'))
if not os.path.isdir(DATA_HOME):
    os.mkdir(DATA_HOME)

document = Document()


def pdf_to_docx(pdf_file, save_home):
    """
    子函数:pdf转为docx
    :param: pdf_file:pdf文件
    :param: save_home:pdf的保存路径
    """
    # rb以二进制读模式打开本地pdf文件
    (file_name, _) = os.path.splitext(pdf_file)
    fn = open(save_home + pdf_file, 'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'  # 添加段落，样式为unordered list类型
                    )
                document.save(save_home + file_name + '.docx')  # 保存这个文档
    return 'pdf 转换 docx 成功'


def doc_docx(doc_name, docx_name):
    """
      子函数:调用word或wps将doc转docx
      :param: docx_name:保存的docx文件名
      :param: doc_name:doc的保存路径
      """
    try:
        # 首先将doc转换成docx
        word = client.Dispatch("kwps.Application")
        doc = word.Documents.Open(doc_name)
        # 使用参数16表示将doc转换成docx
        doc.SaveAs(docx_name, 16)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(e)


def doc_to_docx(name, file_name):
    status = 0
    os.chdir(
        os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir, 'temp/')) + '/' + file_name + '/')
    try:
        output = subprocess.check_output('antiword -mUTF-8 ' + name + '.doc', stderr=subprocess.STDOUT, shell=True)
        out = output.decode('utf8')
        out = str(out).replace("I can't find the name of your HOME directory", "")
        file_word = docx.Document()
        file_word.add_paragraph(out)
        file_word.save(name + ".docx")
    except:
        print("error", name)
        status = 1

    return status


# def word_to_docx(file_home, file_name):
#     """
#        主函数:统一文件格式，统一为docx
#        :param: file_home:文件夹
#        """
#     for file in os.listdir(file_home):
#         if file.find('~$') == -1:
#             if file.find('.doc') != -1 and file.find('.docx') == -1:
#                 name = file[0:file.find('.doc')]
#                 status = doc_to_docx(name, file_name)
#                 if status == 0:
#                     os.remove(file_home + file)
#                     print(file)
#             if file.find('pdf') != -1:
#                 pdf_to_docx(file, file_home)
#                 os.remove(file_home + file)
#                 print(file)
#     return 'doc 转换 docx 成功'
def word_to_docx(file_home):
    """
       主函数:统一文件格式，统一为docx
       :param: file_home:文件夹
       """
    for file in os.listdir(file_home):
        if file.find('.doc') != -1 and file.find('.docx') == -1:
            name = file[0:file.find('.doc')]
            doc_docx(file_home + file, file_home + name + '.docx')
            os.remove(file_home + file)
            print(file)
        if file.find('pdf') != -1:
            pdf_to_docx(file, file_home)
            os.remove(file_home + file)
            print(file)
    for file in os.listdir(file_home):
        if file.find('docx')==-1:
            os.remove(file_home + file)
    return 'doc 转换 docx 成功'


def prepare_docx(content_name, index_name):
    word_to_docx(content_name)
    word_to_docx(index_name)
